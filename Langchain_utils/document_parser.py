"""
文档解析模块

支持 PDF / Word / Excel / TXT 文件的文本提取。
"""
import os
import base64
import logging
from typing import List

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

MAX_TEXT_LENGTH = 10000
ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.txt'}


def parse_document(file_bytes: bytes, filename: str) -> str:
    """
    根据文件扩展名选择解析器提取文本。

    Args:
        file_bytes: 文件二进制数据
        filename: 带扩展名的文件名（如 "报告.pdf"）

    Returns:
        提取的文本内容，不超过 MAX_TEXT_LENGTH 字符
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件格式: {ext}，支持: {', '.join(ALLOWED_EXTENSIONS)}")

    logger.info(f"解析文档: {filename} ({len(file_bytes)} bytes)")

    if ext == '.pdf':
        text = _parse_pdf(file_bytes)
    elif ext == '.docx':
        text = _parse_docx(file_bytes)
    elif ext == '.xlsx':
        text = _parse_xlsx(file_bytes)
    elif ext == '.txt':
        text = _parse_txt(file_bytes)
    else:
        raise ValueError(f"未实现的解析器: {ext}")

    if len(text) > MAX_TEXT_LENGTH:
        logger.info(f"文档文本过长 ({len(text)} 字符)，截断至 {MAX_TEXT_LENGTH}")
        text = text[:MAX_TEXT_LENGTH] + "\n\n[文档内容过长，已截断...]"

    return text


def _parse_pdf(data: bytes) -> str:
    """解析 PDF 文件"""
    try:
        from io import BytesIO
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(data))
        parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                parts.append(page_text)
            if len('\n'.join(parts)) > MAX_TEXT_LENGTH * 2:
                break

        if not parts:
            return "[PDF 文件中未提取到文字内容，可能是扫描件或图片型 PDF]"
        return '\n'.join(parts)
    except Exception as e:
        raise RuntimeError(f"PDF 解析失败: {str(e)}")


def _parse_docx(data: bytes) -> str:
    """解析 Word (.docx) 文件"""
    try:
        from io import BytesIO
        from docx import Document

        doc = Document(BytesIO(data))
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # 同时提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append(' | '.join(cells))

        if not parts:
            return "[Word 文件中未提取到文字内容]"
        return '\n'.join(parts)
    except Exception as e:
        raise RuntimeError(f"Word 解析失败: {str(e)}")


def _parse_xlsx(data: bytes) -> str:
    """解析 Excel 文件，转为 Markdown 表格格式"""
    try:
        from io import BytesIO
        from openpyxl import load_workbook

        wb = load_workbook(BytesIO(data), data_only=True)
        all_parts = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            all_parts.append(f"## 工作表: {sheet_name}\n")

            rows = list(ws.iter_rows(values_only=True))
            if not rows:
                all_parts.append("[空工作表]\n")
                continue

            # 限制行数
            max_rows = min(len(rows), 200)
            rows = rows[:max_rows]

            for i, row in enumerate(rows):
                values = [str(cell) if cell is not None else '' for cell in row]
                all_parts.append('| ' + ' | '.join(values) + ' |')
                if i == 0 and len(rows) > 1:
                    all_parts.append('|' + '|'.join(['---'] * len(values)) + '|')

            all_parts.append('')

        if not all_parts:
            return "[Excel 文件中未提取到数据]"
        return '\n'.join(all_parts)
    except Exception as e:
        raise RuntimeError(f"Excel 解析失败: {str(e)}")


def _parse_txt(data: bytes) -> str:
    """解析纯文本文件"""
    try:
        for encoding in ['utf-8', 'gbk', 'gb2312', 'latin-1']:
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode('utf-8', errors='replace')
    except Exception as e:
        raise RuntimeError(f"文本文件解析失败: {str(e)}")


def chunk_text(text: str, chunk_size: int = 2000, overlap: int = 200) -> List[str]:
    """简单文本分块（供后续检索式注入使用）"""
    if len(text) <= chunk_size:
        return [text]
    chunks = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        chunks.append(text[start:end])
        start += chunk_size - overlap
    return chunks